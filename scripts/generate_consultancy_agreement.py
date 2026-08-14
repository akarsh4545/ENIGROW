#!/usr/bin/env python3
"""Generate concise Enigrow Consultancy Services Agreement (~5–6 pages).

Aligned to a simple Indian consultancy-agreement structure.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    HRFlowable,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

DOWNLOADS = Path.home() / "Downloads"
DOCX_OUT = DOWNLOADS / "Enigrow_Consultancy_Services_Agreement.docx"
PDF_OUT = DOWNLOADS / "Enigrow_Consultancy_Services_Agreement.pdf"


# ═════════════════════════════════════════════════════════════
# WORD HELPERS
# ═════════════════════════════════════════════════════════════

def set_run_font(run, size_pt=11, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_pf(
    paragraph,
    *,
    space_before=0,
    space_after=5,
    line_spacing=1.2,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    keep_with_next=False,
    page_break_before=False,
    left_indent=None,
):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line_spacing
    fmt.alignment = align
    fmt.widow_control = True
    fmt.keep_with_next = keep_with_next
    fmt.page_break_before = page_break_before
    if left_indent is not None:
        fmt.left_indent = left_indent


def add_text(
    doc,
    text,
    *,
    size=11,
    bold=False,
    italic=False,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    space_before=0,
    space_after=5,
    keep_with_next=False,
    page_break_before=False,
    left_indent=None,
):
    p = doc.add_paragraph()
    set_pf(
        p,
        space_before=space_before,
        space_after=space_after,
        align=align,
        keep_with_next=keep_with_next,
        page_break_before=page_break_before,
        left_indent=left_indent,
    )
    parts, remaining = [], text
    while "<b>" in remaining and "</b>" in remaining:
        before, rest = remaining.split("<b>", 1)
        bold_text, remaining = rest.split("</b>", 1)
        if before:
            parts.append((before, False))
        parts.append((bold_text, True))
    if remaining:
        parts.append((remaining, False))
    if not parts:
        parts = [(text, bold)]
    for chunk, is_bold in parts:
        if not chunk:
            continue
        run = p.add_run(chunk.replace("&amp;", "&"))
        set_run_font(run, size_pt=size, bold=is_bold or bold, italic=italic)
    return p


def H(doc, text, *, page_break_before=False):
    return add_text(
        doc,
        text,
        size=12,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        space_before=9,
        space_after=3,
        keep_with_next=True,
        page_break_before=page_break_before,
    )


def B(doc, text, *, keep_with_next=False, left_indent=None, space_after=5):
    return add_text(
        doc,
        text,
        size=11,
        space_after=space_after,
        keep_with_next=keep_with_next,
        left_indent=left_indent,
    )


def C(doc, text):
    return add_text(
        doc,
        text,
        size=11,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=5,
        space_after=5,
    )


def hr(doc):
    p = doc.add_paragraph()
    set_pf(p, space_before=1, space_after=5, align=WD_ALIGN_PARAGRAPH.CENTER)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "222222")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_field(paragraph, instr):
    run = paragraph.add_run()
    set_run_font(run, size_pt=8)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    i = OxmlElement("w:instrText")
    i.set(qn("xml:space"), "preserve")
    i.text = instr
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, i, sep, end])


def build_docx() -> Path:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.7)
        section.bottom_margin = Cm(1.7)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    add_text(
        doc,
        "CONSULTANCY SERVICES AGREEMENT",
        size=14,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=3,
        keep_with_next=True,
    )
    hr(doc)

    B(
        doc,
        'This Consultancy Services Agreement ("Agreement") is made and entered into on this ___ day of ______________, 20__ ("Effective Date"). Where this Agreement is executed by the Parties on different dates, the later date of execution shall be deemed to be the Effective Date.',
    )
    C(doc, "BY AND BETWEEN")
    B(
        doc,
        '<b>ENIGROW STARTUP ADVISORY PRIVATE LIMITED</b>, a company incorporated under the Companies Act, 2013, having its registered office at B-128, 1st Floor, Sector-2, Noida, Gautam Buddha Nagar, Uttar Pradesh – 201301, CIN: U82990UW2026PTC255445, Email: support@enigrow.co.in, Website: www.enigrow.co.in (hereinafter the "<b>Service Provider</b>" or "<b>Enigrow</b>");',
    )
    C(doc, "AND")
    B(
        doc,
        '<b>[CLIENT NAME]</b>, a [Proprietorship / Partnership / LLP / Private Limited Company / Individual], having its registered office / principal place of business at [CLIENT ADDRESS], PAN: [●], GSTIN (if applicable): [●], Email: [●], Mobile: [●], acting through its Authorized Signatory [NAME], Designation: [●] (hereinafter the "<b>Client</b>").',
    )
    B(
        doc,
        'The Service Provider and the Client are hereinafter individually a "<b>Party</b>" and collectively the "<b>Parties</b>".',
    )
    B(
        doc,
        "WHEREAS the Service Provider provides consultancy and advisory services and the Client desires to avail such services on the terms set out herein, the Parties agree as follows:",
    )

    H(doc, "1. PURPOSE")
    B(
        doc,
        "The Client engages the Service Provider to provide consultancy and advisory assistance relating to funding, government schemes, loan applications, documentation, application preparation, and business advisory support, as more particularly described in Annexure A. Nothing in this Agreement promises any specific funding or approval outcome.",
    )

    H(doc, "2. DEFINITIONS")
    B(
        doc,
        'In this Agreement: <b>"Agreement"</b> means this Consultancy Services Agreement including its Annexures; <b>"Service Provider"</b> means Enigrow Startup Advisory Private Limited; <b>"Client"</b> means the person or entity named above; <b>"Services"</b> means the consultancy and advisory services described in this Agreement and Annexure A; <b>"Consultancy Fee"</b> means the professional charges payable for the Services as set out in Annexure B; and <b>"Success Fee"</b> means any fee payable upon the successful event specified in Annexure B.',
    )

    H(doc, "3. SCOPE OF SERVICES")
    B(
        doc,
        "3.1 Subject to Annexure A and with reasonable professional care, the Service Provider shall provide the agreed Services, which may include: (a) eligibility and documentation assistance; (b) preparation and review of required documents; (c) DPR, business plan, and financial projections where applicable; (d) application preparation; (e) submission assistance where authorized by the Client; (f) coordination and follow-up with relevant banks, institutions, or authorities; and (g) advisory support.",
    )
    B(
        doc,
        "3.2 The Service Provider does not control approval, sanction, disbursement, government or authority decisions, bank or institution decisions, policy changes, or third-party processing timelines. Outcomes of that nature remain subject to Clause 6.",
    )

    H(doc, "4. CLIENT OBLIGATIONS")
    B(
        doc,
        "4.1 The Client shall: (a) provide complete, accurate, and genuine information and documents within the required timeline; (b) provide signatures, declarations, authorizations, and clarifications when required; (c) review information and documents prepared on its behalf; (d) cooperate with the Service Provider; and (e) make payments on time in accordance with Clause 5 and Annexure B.",
    )
    B(
        doc,
        "4.2 Any delay, rejection, deficiency, or additional requirement arising from incomplete, inaccurate, misleading, outdated, or delayed information or documents provided by the Client shall be the Client's responsibility. The Service Provider is not responsible for information originating from the Client.",
    )
    B(
        doc,
        "4.3 Timelines communicated by the Service Provider are estimates only and may be extended where delay results from the Client, a bank, government authority, financial institution, third-party service provider, policy or eligibility changes, or circumstances outside the Service Provider's reasonable control.",
    )

    H(doc, "5. CONSULTANCY FEE AND PAYMENT")
    B(
        doc,
        "5.1 The fees payable are as set out in Annexure B. Agreed initial / advance fees become payable according to Annexure B. The Service Provider may commence substantive work after receipt of the required initial payment and documents. Payment obligations are independent of the ultimate approval, sanction, or disbursement decision.",
    )
    B(
        doc,
        "5.2 Where a Success Fee applies, it becomes payable upon the successful event specified in Annexure B (such as sanction, disbursement, receipt of funding, receipt of scheme benefit, or another expressly agreed event). Once that trigger occurs, the Success Fee becomes due irrespective of whether the Client subsequently uses, retains, returns, restructures, or otherwise deals with the sanctioned or disbursed amount or benefit.",
    )
    B(
        doc,
        "5.3 Unless expressly agreed otherwise in writing, fees for Services already commenced or provided are non-refundable. Payments shall be made by bank transfer to the account stated in Annexure B, or by any other mutually agreed method.",
    )

    H(doc, "6. NO GUARANTEE")
    B(
        doc,
        "The Service Provider does not guarantee eligibility, approval, sanction, funding, disbursement, subsidy or benefit, investor acceptance, or any specific processing time. Such decisions remain solely with the relevant bank, authority, institution, investor, or other third party, and criteria, policies, and timelines may change.",
    )

    H(doc, "7. CONFIDENTIALITY")
    B(
        doc,
        "Each Party shall keep confidential the other Party's non-public documents, financial information, business information, personal or business information, and application-related information, and shall not disclose the same except as reasonably necessary for providing the Services or as required by law or competent authority.",
    )

    H(doc, "8. LIABILITY")
    B(
        doc,
        "The Service Provider shall exercise reasonable professional care. The Service Provider shall not be responsible for rejection by a bank or authority, policy changes, third-party decisions, delays outside its reasonable control, or inaccurate information supplied by the Client. To the maximum extent permitted by law, the Service Provider's total liability under this Agreement shall not exceed the Consultancy Fees actually received under the engagement.",
    )

    H(doc, "9. TERM")
    B(
        doc,
        "This Agreement shall remain valid for one (1) year from the Effective Date unless terminated earlier in accordance with Clause 10.",
    )

    H(doc, "10. TERMINATION")
    B(
        doc,
        "Either Party may terminate this Agreement by giving thirty (30) days' prior written notice. The Service Provider may suspend or terminate the Services with immediate effect in case of serious breach, non-payment, false or forged documents, or unlawful activity. Termination does not cancel fees already accrued or payment obligations for Services already provided, does not require the Service Provider to refund fees merely because the Client chooses to terminate, and does not affect provisions intended to survive termination.",
    )

    H(doc, "11. INDEPENDENT RELATIONSHIP")
    B(
        doc,
        "The Parties are independent contractors. Nothing in this Agreement creates an employment, partnership, joint venture, or agency relationship. Neither Party shall have authority to create any obligation on behalf of the other.",
    )

    H(doc, "12. THIRD-PARTY RIGHTS")
    B(
        doc,
        "No person other than the Parties shall have any right to enforce any term of this Agreement.",
    )

    H(doc, "13. MODIFICATION")
    B(
        doc,
        "This Agreement may be amended or modified only by a written instrument signed by authorized representatives of both Parties.",
    )

    H(doc, "14. SEVERABILITY")
    B(
        doc,
        "If any provision of this Agreement is held invalid or unenforceable, the remaining provisions shall continue in full force and effect.",
    )

    H(doc, "15. ENFORCEMENT AND WAIVER")
    B(
        doc,
        "No failure or delay by either Party in exercising any right or remedy under this Agreement shall constitute a waiver of that right or remedy.",
    )

    H(doc, "16. NOTICES")
    B(
        doc,
        "Notices under this Agreement may be given by email, registered post, courier, or other written communication to the addresses or contact details stated in this Agreement.",
    )

    H(doc, "17. ENTIRE AGREEMENT")
    B(
        doc,
        "This Agreement constitutes the entire understanding between the Parties and supersedes all prior discussions and understandings relating to its subject matter.",
    )

    H(doc, "18. COUNTERPARTS")
    B(
        doc,
        "This Agreement may be executed in counterparts and by electronic means where legally permissible, each of which shall be deemed an original.",
    )

    H(doc, "19. GOVERNING LAW AND ARBITRATION")
    B(
        doc,
        "19.1 This Agreement shall be governed by the laws of India. Subject to arbitration, the courts at Noida, Gautam Buddha Nagar, Uttar Pradesh shall have jurisdiction.",
    )
    B(
        doc,
        "19.2 The Parties shall first attempt amicable resolution. If unresolved within thirty (30) days, the dispute shall be referred to arbitration under the Arbitration and Conciliation Act, 1996, before a sole arbitrator mutually appointed by the Parties. If the Parties cannot agree on the arbitrator, appointment shall be made in accordance with applicable law. The seat and venue shall be Noida, Uttar Pradesh, the language shall be English, and the award shall be final and binding.",
    )

    H(doc, "20. CLIENT UNDERTAKING")
    B(
        doc,
        "The Client confirms that: (a) the information and documents provided are genuine and accurate; (b) it will comply with its obligations and provide required cooperation; (c) it understands that approval or disbursement is not guaranteed, as stated in Clause 6; and (d) it has read and accepted this Agreement. Material false, forged, or misleading information or documents may result in suspension or termination of Services, with the Client remaining liable for accrued fees.",
    )

    B(
        doc,
        "<b>IN WITNESS WHEREOF</b>, the Parties have executed this Agreement on the date first written above.",
    )

    add_text(
        doc,
        "FOR THE SERVICE PROVIDER",
        size=11,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        space_before=10,
        space_after=2,
        keep_with_next=True,
    )
    add_text(
        doc,
        "ENIGROW STARTUP ADVISORY PRIVATE LIMITED",
        size=11,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        space_after=6,
        keep_with_next=True,
    )
    for line in [
        "Authorized Signatory: ______________________________",
        "Name: ____________________________________________",
        "Designation: _______________________________________",
        "Place: Noida, Uttar Pradesh",
        "Date: _____________________________________________",
        "Company Seal: _____________________________________",
    ]:
        add_text(doc, line, size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)

    add_text(
        doc,
        "FOR THE CLIENT",
        size=11,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        space_before=10,
        space_after=2,
        keep_with_next=True,
    )
    add_text(
        doc,
        "[CLIENT NAME]",
        size=11,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        space_after=6,
        keep_with_next=True,
    )
    for line in [
        "Authorized Signatory: ______________________________",
        "Name: ____________________________________________",
        "Designation: _______________________________________",
        "Place: _____________________________________________",
        "Date: _____________________________________________",
        "Company Seal (if applicable): _______________________",
    ]:
        add_text(doc, line, size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)

    # Annexure A
    H(doc, "ANNEXURE A — SERVICES & DELIVERABLES", page_break_before=True)
    hr(doc)
    B(doc, "This Annexure forms an integral part of the Agreement.")
    B(doc, "<b>Service:</b> _______________________________________________")
    B(
        doc,
        "<b>Scope</b> (as applicable): eligibility/documentation assistance; application preparation; DPR/business plan/projections where applicable; submission assistance; coordination/follow-up; and advisory support.",
    )
    B(
        doc,
        "<b>Key Deliverables</b> (as applicable): DPR / Pitch Deck / Projection Report; business plan and financial projections; application forms and supporting documents; and submission/acknowledgment records where applicable.",
    )
    B(doc, "<b>Target Scheme / Product:</b> _________________________________")
    B(doc, "<b>Estimated Timeline:</b> ______________________________________")
    B(doc, "<b>Special Instructions:</b> _____________________________________")
    add_text(
        doc,
        "Note: Subject to Clause 6 (No Guarantee).",
        size=10,
        italic=True,
        space_before=4,
    )

    # Annexure B — continues after Annexure A (no forced page break)
    H(doc, "ANNEXURE B — FEE SCHEDULE")
    hr(doc)
    B(
        doc,
        "This Annexure forms an integral part of the Agreement. Amounts are in INR and exclusive of applicable taxes unless otherwise stated.",
    )

    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Particulars", "Amount / Terms"),
        ("Advance Consultancy Fee", "INR ______________"),
        ("Documentation / Professional Charges", "INR ______________"),
        ("Success Fee (if applicable)", "INR ______________ / ____%"),
        ("Other Charges", "INR ______________"),
        ("Payment Milestones", "________________________"),
    ]
    for i, (left, right) in enumerate(rows):
        for cell, val, header in (
            (table.rows[i].cells[0], left, i == 0),
            (table.rows[i].cells[1], right, i == 0),
        ):
            cell.text = ""
            p = cell.paragraphs[0]
            set_pf(p, space_before=3, space_after=3, align=WD_ALIGN_PARAGRAPH.LEFT)
            run = p.add_run(val)
            set_run_font(run, size_pt=10.5, bold=header)
            if i == 0:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "EFEFEF")
                shd.set(qn("w:val"), "clear")
                tcPr.append(shd)

    add_text(doc, " ", size=11, space_after=6, align=WD_ALIGN_PARAGRAPH.LEFT)
    B(doc, "<b>Bank Account Details of the Service Provider</b>", space_after=3)
    for line in [
        "Account Name: ENIGROW STARTUP ADVISORY PRIVATE LIMITED",
        "Bank: Kotak Mahindra Bank",
        "Account Number: 6151522453",
        "IFSC: KKBK0000201",
        "Branch: Delhi – Nehru Place",
    ]:
        add_text(doc, line, size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=1)

    # Annexure C — compact numbered list (continues on same page)
    H(doc, "ANNEXURE C — DOCUMENT CHECKLIST")
    hr(doc)
    B(
        doc,
        "The Client shall provide the following documents as applicable. Additional documents may be requested depending on the requirements of the relevant authority, bank, financial institution, or scheme.",
    )
    checklist = [
        "PAN Card of Client / Promoters / Directors",
        "Aadhaar / Passport / other KYC documents",
        "Proof of address",
        "GST Registration (if applicable)",
        "Udyam / MSME Registration (if applicable)",
        "Incorporation / Partnership / LLP / Shop Act documents (as applicable)",
        "ITR for last 2–3 years",
        "Financial statements for last 2–3 years",
        "Bank statements for last 6–12 months",
        "Existing loan documents (if any)",
        "Project quotations / estimates / invoices (if applicable)",
        "Any other document as may be required",
    ]
    chk = doc.add_table(rows=6, cols=2)
    for i in range(6):
        left = f"{i + 1}.  ☐  {checklist[i]}"
        right = f"{i + 7}.  ☐  {checklist[i + 6]}"
        for col, val in ((0, left), (1, right)):
            cell = chk.rows[i].cells[col]
            cell.text = ""
            p = cell.paragraphs[0]
            set_pf(p, space_before=1, space_after=1, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_run_font(p.add_run(val), size_pt=10)

    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(
            hp.add_run("Enigrow Startup Advisory Pvt. Ltd.  |  Consultancy Services Agreement"),
            size_pt=8,
        )
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(fp.add_run("Page "), size_pt=8)
        add_field(fp, " PAGE ")
        set_run_font(fp.add_run(" of "), size_pt=8)
        add_field(fp, " NUMPAGES ")

    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(DOCX_OUT))
    return DOCX_OUT


# ═════════════════════════════════════════════════════════════
# PDF
# ═════════════════════════════════════════════════════════════

def styles():
    leading = 13.5
    return {
        "title": ParagraphStyle(
            "T", fontName="Times-Bold", fontSize=13, leading=16, alignment=TA_CENTER, spaceAfter=3
        ),
        "section": ParagraphStyle(
            "S", fontName="Times-Bold", fontSize=11, leading=13, spaceBefore=8, spaceAfter=2
        ),
        "body": ParagraphStyle(
            "B", fontName="Times-Roman", fontSize=10.5, leading=leading, alignment=TA_JUSTIFY, spaceAfter=3.5
        ),
        "center": ParagraphStyle(
            "C", fontName="Times-Bold", fontSize=10.5, leading=leading, alignment=TA_CENTER, spaceBefore=4, spaceAfter=4
        ),
        "bullet": ParagraphStyle(
            "L", fontName="Times-Roman", fontSize=10.5, leading=leading, leftIndent=12, spaceAfter=1.5
        ),
        "sig": ParagraphStyle(
            "G", fontName="Times-Roman", fontSize=10, leading=12.5, spaceAfter=1.5
        ),
        "sig_label": ParagraphStyle(
            "GL", fontName="Times-Bold", fontSize=10.5, leading=13, spaceBefore=7, spaceAfter=2
        ),
        "note": ParagraphStyle(
            "N", fontName="Times-Italic", fontSize=9.5, leading=12, spaceBefore=3, spaceAfter=3
        ),
    }


def P(s, text, style="body"):
    return Paragraph(text, s[style])


def decorate(total):
    def _draw(canvas, doc):
        canvas.saveState()
        canvas.setFont("Times-Roman", 8)
        canvas.setFillColorRGB(0.25, 0.25, 0.25)
        canvas.drawCentredString(
            A4[0] / 2,
            A4[1] - 0.38 * inch,
            "Enigrow Startup Advisory Pvt. Ltd.  |  Consultancy Services Agreement",
        )
        canvas.drawCentredString(A4[0] / 2, 0.38 * inch, f"Page {canvas.getPageNumber()} of {total}")
        canvas.restoreState()

    return _draw


def build_story():
    s = styles()
    story = []

    story.append(P(s, "CONSULTANCY SERVICES AGREEMENT", "title"))
    story.append(HRFlowable(width="100%", thickness=1, color="#222222", spaceAfter=5))
    story.append(
        P(
            s,
            'This Consultancy Services Agreement ("Agreement") is made and entered into on this ___ day of ______________, 20__ ("Effective Date"). Where this Agreement is executed by the Parties on different dates, the later date of execution shall be deemed to be the Effective Date.',
        )
    )
    story.append(P(s, "BY AND BETWEEN", "center"))
    story.append(
        P(
            s,
            '<b>ENIGROW STARTUP ADVISORY PRIVATE LIMITED</b>, a company incorporated under the Companies Act, 2013, having its registered office at B-128, 1st Floor, Sector-2, Noida, Gautam Buddha Nagar, Uttar Pradesh – 201301, CIN: U82990UW2026PTC255445, Email: support@enigrow.co.in, Website: www.enigrow.co.in (hereinafter the "<b>Service Provider</b>" or "<b>Enigrow</b>");',
        )
    )
    story.append(P(s, "AND", "center"))
    story.append(
        P(
            s,
            '<b>[CLIENT NAME]</b>, a [Proprietorship / Partnership / LLP / Private Limited Company / Individual], having its registered office / principal place of business at [CLIENT ADDRESS], PAN: [●], GSTIN (if applicable): [●], Email: [●], Mobile: [●], acting through its Authorized Signatory [NAME], Designation: [●] (hereinafter the "<b>Client</b>").',
        )
    )
    story.append(
        P(
            s,
            'The Service Provider and the Client are hereinafter individually a "<b>Party</b>" and collectively the "<b>Parties</b>".',
        )
    )
    story.append(
        P(
            s,
            "WHEREAS the Service Provider provides consultancy and advisory services and the Client desires to avail such services on the terms set out herein, the Parties agree as follows:",
        )
    )

    story.append(P(s, "1. PURPOSE", "section"))
    story.append(
        P(
            s,
            "The Client engages the Service Provider to provide consultancy and advisory assistance relating to funding, government schemes, loan applications, documentation, application preparation, and business advisory support, as more particularly described in Annexure A. Nothing in this Agreement promises any specific funding or approval outcome.",
        )
    )

    story.append(P(s, "2. DEFINITIONS", "section"))
    story.append(
        P(
            s,
            'In this Agreement: <b>"Agreement"</b> means this Consultancy Services Agreement including its Annexures; <b>"Service Provider"</b> means Enigrow Startup Advisory Private Limited; <b>"Client"</b> means the person or entity named above; <b>"Services"</b> means the consultancy and advisory services described in this Agreement and Annexure A; <b>"Consultancy Fee"</b> means the professional charges payable for the Services as set out in Annexure B; and <b>"Success Fee"</b> means any fee payable upon the successful event specified in Annexure B.',
        )
    )

    story.append(P(s, "3. SCOPE OF SERVICES", "section"))
    story.append(
        P(
            s,
            "3.1 Subject to Annexure A and with reasonable professional care, the Service Provider shall provide the agreed Services, which may include: (a) eligibility and documentation assistance; (b) preparation and review of required documents; (c) DPR, business plan, and financial projections where applicable; (d) application preparation; (e) submission assistance where authorized by the Client; (f) coordination and follow-up with relevant banks, institutions, or authorities; and (g) advisory support.",
        )
    )
    story.append(
        P(
            s,
            "3.2 The Service Provider does not control approval, sanction, disbursement, government or authority decisions, bank or institution decisions, policy changes, or third-party processing timelines. Outcomes of that nature remain subject to Clause 6.",
        )
    )

    story.append(P(s, "4. CLIENT OBLIGATIONS", "section"))
    story.append(
        P(
            s,
            "4.1 The Client shall: (a) provide complete, accurate, and genuine information and documents within the required timeline; (b) provide signatures, declarations, authorizations, and clarifications when required; (c) review information and documents prepared on its behalf; (d) cooperate with the Service Provider; and (e) make payments on time in accordance with Clause 5 and Annexure B.",
        )
    )
    story.append(
        P(
            s,
            "4.2 Any delay, rejection, deficiency, or additional requirement arising from incomplete, inaccurate, misleading, outdated, or delayed information or documents provided by the Client shall be the Client's responsibility. The Service Provider is not responsible for information originating from the Client.",
        )
    )
    story.append(
        P(
            s,
            "4.3 Timelines communicated by the Service Provider are estimates only and may be extended where delay results from the Client, a bank, government authority, financial institution, third-party service provider, policy or eligibility changes, or circumstances outside the Service Provider's reasonable control.",
        )
    )

    story.append(P(s, "5. CONSULTANCY FEE AND PAYMENT", "section"))
    story.append(
        P(
            s,
            "5.1 The fees payable are as set out in Annexure B. Agreed initial / advance fees become payable according to Annexure B. The Service Provider may commence substantive work after receipt of the required initial payment and documents. Payment obligations are independent of the ultimate approval, sanction, or disbursement decision.",
        )
    )
    story.append(
        P(
            s,
            "5.2 Where a Success Fee applies, it becomes payable upon the successful event specified in Annexure B (such as sanction, disbursement, receipt of funding, receipt of scheme benefit, or another expressly agreed event). Once that trigger occurs, the Success Fee becomes due irrespective of whether the Client subsequently uses, retains, returns, restructures, or otherwise deals with the sanctioned or disbursed amount or benefit.",
        )
    )
    story.append(
        P(
            s,
            "5.3 Unless expressly agreed otherwise in writing, fees for Services already commenced or provided are non-refundable. Payments shall be made by bank transfer to the account stated in Annexure B, or by any other mutually agreed method.",
        )
    )

    story.append(P(s, "6. NO GUARANTEE", "section"))
    story.append(
        P(
            s,
            "The Service Provider does not guarantee eligibility, approval, sanction, funding, disbursement, subsidy or benefit, investor acceptance, or any specific processing time. Such decisions remain solely with the relevant bank, authority, institution, investor, or other third party, and criteria, policies, and timelines may change.",
        )
    )

    story.append(P(s, "7. CONFIDENTIALITY", "section"))
    story.append(
        P(
            s,
            "Each Party shall keep confidential the other Party's non-public documents, financial information, business information, personal or business information, and application-related information, and shall not disclose the same except as reasonably necessary for providing the Services or as required by law or competent authority.",
        )
    )

    story.append(P(s, "8. LIABILITY", "section"))
    story.append(
        P(
            s,
            "The Service Provider shall exercise reasonable professional care. The Service Provider shall not be responsible for rejection by a bank or authority, policy changes, third-party decisions, delays outside its reasonable control, or inaccurate information supplied by the Client. To the maximum extent permitted by law, the Service Provider's total liability under this Agreement shall not exceed the Consultancy Fees actually received under the engagement.",
        )
    )

    story.append(P(s, "9. TERM", "section"))
    story.append(
        P(
            s,
            "This Agreement shall remain valid for one (1) year from the Effective Date unless terminated earlier in accordance with Clause 10.",
        )
    )

    story.append(P(s, "10. TERMINATION", "section"))
    story.append(
        P(
            s,
            "Either Party may terminate this Agreement by giving thirty (30) days' prior written notice. The Service Provider may suspend or terminate the Services with immediate effect in case of serious breach, non-payment, false or forged documents, or unlawful activity. Termination does not cancel fees already accrued or payment obligations for Services already provided, does not require the Service Provider to refund fees merely because the Client chooses to terminate, and does not affect provisions intended to survive termination.",
        )
    )

    story.append(P(s, "11. INDEPENDENT RELATIONSHIP", "section"))
    story.append(
        P(
            s,
            "The Parties are independent contractors. Nothing in this Agreement creates an employment, partnership, joint venture, or agency relationship. Neither Party shall have authority to create any obligation on behalf of the other.",
        )
    )

    story.append(P(s, "12. THIRD-PARTY RIGHTS", "section"))
    story.append(
        P(
            s,
            "No person other than the Parties shall have any right to enforce any term of this Agreement.",
        )
    )

    story.append(P(s, "13. MODIFICATION", "section"))
    story.append(
        P(
            s,
            "This Agreement may be amended or modified only by a written instrument signed by authorized representatives of both Parties.",
        )
    )

    story.append(P(s, "14. SEVERABILITY", "section"))
    story.append(
        P(
            s,
            "If any provision of this Agreement is held invalid or unenforceable, the remaining provisions shall continue in full force and effect.",
        )
    )

    story.append(P(s, "15. ENFORCEMENT AND WAIVER", "section"))
    story.append(
        P(
            s,
            "No failure or delay by either Party in exercising any right or remedy under this Agreement shall constitute a waiver of that right or remedy.",
        )
    )

    story.append(P(s, "16. NOTICES", "section"))
    story.append(
        P(
            s,
            "Notices under this Agreement may be given by email, registered post, courier, or other written communication to the addresses or contact details stated in this Agreement.",
        )
    )

    story.append(P(s, "17. ENTIRE AGREEMENT", "section"))
    story.append(
        P(
            s,
            "This Agreement constitutes the entire understanding between the Parties and supersedes all prior discussions and understandings relating to its subject matter.",
        )
    )

    story.append(P(s, "18. COUNTERPARTS", "section"))
    story.append(
        P(
            s,
            "This Agreement may be executed in counterparts and by electronic means where legally permissible, each of which shall be deemed an original.",
        )
    )

    story.append(P(s, "19. GOVERNING LAW AND ARBITRATION", "section"))
    story.append(
        P(
            s,
            "19.1 This Agreement shall be governed by the laws of India. Subject to arbitration, the courts at Noida, Gautam Buddha Nagar, Uttar Pradesh shall have jurisdiction.",
        )
    )
    story.append(
        P(
            s,
            "19.2 The Parties shall first attempt amicable resolution. If unresolved within thirty (30) days, the dispute shall be referred to arbitration under the Arbitration and Conciliation Act, 1996, before a sole arbitrator mutually appointed by the Parties. If the Parties cannot agree on the arbitrator, appointment shall be made in accordance with applicable law. The seat and venue shall be Noida, Uttar Pradesh, the language shall be English, and the award shall be final and binding.",
        )
    )

    story.append(P(s, "20. CLIENT UNDERTAKING", "section"))
    story.append(
        P(
            s,
            "The Client confirms that: (a) the information and documents provided are genuine and accurate; (b) it will comply with its obligations and provide required cooperation; (c) it understands that approval or disbursement is not guaranteed, as stated in Clause 6; and (d) it has read and accepted this Agreement. Material false, forged, or misleading information or documents may result in suspension or termination of Services, with the Client remaining liable for accrued fees.",
        )
    )

    story.append(
        P(
            s,
            "<b>IN WITNESS WHEREOF</b>, the Parties have executed this Agreement on the date first written above.",
        )
    )
    story.append(Spacer(1, 6))
    story.append(P(s, "FOR THE SERVICE PROVIDER", "sig_label"))
    story.append(P(s, "ENIGROW STARTUP ADVISORY PRIVATE LIMITED", "sig"))
    story.append(Spacer(1, 4))
    for line in [
        "Authorized Signatory: ______________________________",
        "Name: ____________________________________________",
        "Designation: _______________________________________",
        "Place: Noida, Uttar Pradesh",
        "Date: _____________________________________________",
        "Company Seal: _____________________________________",
    ]:
        story.append(P(s, line, "sig"))
    story.append(Spacer(1, 6))
    story.append(P(s, "FOR THE CLIENT", "sig_label"))
    story.append(P(s, "[CLIENT NAME]", "sig"))
    story.append(Spacer(1, 4))
    for line in [
        "Authorized Signatory: ______________________________",
        "Name: ____________________________________________",
        "Designation: _______________________________________",
        "Place: _____________________________________________",
        "Date: _____________________________________________",
        "Company Seal (if applicable): _______________________",
    ]:
        story.append(P(s, line, "sig"))

    # Annexure A
    story.append(PageBreak())
    story.append(P(s, "ANNEXURE A — SERVICES & DELIVERABLES", "title"))
    story.append(HRFlowable(width="100%", thickness=0.8, color="#222222", spaceAfter=5))
    story.append(P(s, "This Annexure forms an integral part of the Agreement."))
    story.append(P(s, "<b>Service:</b> _______________________________________________"))
    story.append(
        P(
            s,
            "<b>Scope</b> (as applicable): eligibility/documentation assistance; application preparation; DPR/business plan/projections where applicable; submission assistance; coordination/follow-up; and advisory support.",
        )
    )
    story.append(
        P(
            s,
            "<b>Key Deliverables</b> (as applicable): DPR / Pitch Deck / Projection Report; business plan and financial projections; application forms and supporting documents; and submission/acknowledgment records where applicable.",
        )
    )
    story.append(P(s, "<b>Target Scheme / Product:</b> _________________________________"))
    story.append(P(s, "<b>Estimated Timeline:</b> ______________________________________"))
    story.append(P(s, "<b>Special Instructions:</b> _____________________________________"))
    story.append(P(s, "Note: Subject to Clause 6 (No Guarantee).", "note"))

    # Annexure B + C continue after Annexure A (no forced page break)
    story.append(Spacer(1, 8))
    story.append(P(s, "ANNEXURE B — FEE SCHEDULE", "title"))
    story.append(HRFlowable(width="100%", thickness=0.8, color="#222222", spaceAfter=4))
    story.append(
        P(
            s,
            "This Annexure forms an integral part of the Agreement. Amounts are in INR and exclusive of applicable taxes unless otherwise stated.",
        )
    )
    story.append(Spacer(1, 3))
    cell = ParagraphStyle("Cell", fontName="Times-Roman", fontSize=10, leading=12)
    cell_b = ParagraphStyle("CellB", fontName="Times-Bold", fontSize=10, leading=12)
    fee_data = [
        [Paragraph("<b>Particulars</b>", cell_b), Paragraph("<b>Amount / Terms</b>", cell_b)],
        [Paragraph("Advance Consultancy Fee", cell), Paragraph("INR ______________", cell)],
        [Paragraph("Documentation / Professional Charges", cell), Paragraph("INR ______________", cell)],
        [Paragraph("Success Fee (if applicable)", cell), Paragraph("INR ______________ / ____%", cell)],
        [Paragraph("Other Charges", cell), Paragraph("INR ______________", cell)],
        [Paragraph("Payment Milestones", cell), Paragraph("________________________", cell)],
    ]
    table = Table(fee_data, colWidths=[3.2 * inch, 3.6 * inch], rowHeights=[20, 22, 22, 22, 22, 22])
    table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#444444")),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#efefef")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]
        )
    )
    story.append(table)
    story.append(Spacer(1, 5))
    story.append(P(s, "<b>Bank Account Details of the Service Provider</b>"))
    for line in [
        "Account Name: ENIGROW STARTUP ADVISORY PRIVATE LIMITED",
        "Bank: Kotak Mahindra Bank",
        "Account Number: 6151522453",
        "IFSC: KKBK0000201",
        "Branch: Delhi – Nehru Place",
    ]:
        story.append(P(s, line, "sig"))

    # Annexure C — keep with fee schedule to avoid orphan checklist page
    checklist = [
        "PAN Card of Client / Promoters / Directors",
        "Aadhaar / Passport / other KYC documents",
        "Proof of address",
        "GST Registration (if applicable)",
        "Udyam / MSME Registration (if applicable)",
        "Incorporation / Partnership / LLP / Shop Act documents (as applicable)",
        "ITR for last 2–3 years",
        "Financial statements for last 2–3 years",
        "Bank statements for last 6–12 months",
        "Existing loan documents (if any)",
        "Project quotations / estimates / invoices (if applicable)",
        "Any other document as may be required",
    ]
    checklist_flow = [
        Spacer(1, 6),
        P(s, "ANNEXURE C — DOCUMENT CHECKLIST", "section"),
        HRFlowable(width="100%", thickness=0.8, color="#222222", spaceAfter=3),
        P(
            s,
            "The Client shall provide the following documents as applicable. Additional documents may be requested depending on the requirements of the relevant authority, bank, financial institution, or scheme.",
        ),
    ]
    # Two-column compact checklist
    left = checklist[:6]
    right = checklist[6:]
    rows = []
    for i in range(6):
        l = f"{i + 1}.  ☐  {left[i]}"
        r = f"{i + 7}.  ☐  {right[i]}" if i < len(right) else ""
        rows.append([Paragraph(l, s["bullet"]), Paragraph(r, s["bullet"])])
    chk = Table(rows, colWidths=[3.45 * inch, 3.45 * inch])
    chk.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 1),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
            ]
        )
    )
    checklist_flow.append(chk)
    story.append(KeepTogether(checklist_flow))

    return story


def build_pdf() -> Path:
    PDF_OUT.parent.mkdir(parents=True, exist_ok=True)
    margins = dict(
        leftMargin=0.8 * inch,
        rightMargin=0.8 * inch,
        topMargin=0.65 * inch,
        bottomMargin=0.6 * inch,
    )
    page_count = [0]

    def _count(canvas, doc):
        page_count[0] = max(page_count[0], canvas.getPageNumber())

    tmp = PDF_OUT.with_suffix(".tmp.pdf")
    SimpleDocTemplate(
        str(tmp),
        pagesize=A4,
        title="Consultancy Services Agreement — Enigrow Startup Advisory Private Limited",
        author="Enigrow Startup Advisory Private Limited",
        **margins,
    ).build(build_story(), onFirstPage=_count, onLaterPages=_count)

    SimpleDocTemplate(
        str(PDF_OUT),
        pagesize=A4,
        title="Consultancy Services Agreement — Enigrow Startup Advisory Private Limited",
        author="Enigrow Startup Advisory Private Limited",
        **margins,
    ).build(build_story(), onFirstPage=decorate(page_count[0]), onLaterPages=decorate(page_count[0]))
    if tmp.exists():
        tmp.unlink()
    return PDF_OUT


def main():
    docx = build_docx()
    pdf = build_pdf()
    from pypdf import PdfReader

    pages = len(PdfReader(str(pdf)).pages)
    print(f"Saved DOCX: {docx}")
    print(f"Saved PDF:  {pdf}")
    print(f"PDF pages:  {pages}")


if __name__ == "__main__":
    main()
